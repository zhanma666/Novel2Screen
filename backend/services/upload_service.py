import os
import hashlib
import re
from fastapi import UploadFile
from sqlalchemy.orm import Session
from models import Project, SourceDocument, Chapter
from utils import generate_id
from exceptions import NotFoundError, BadRequestError
from config import get_settings
from logger import logger

settings = get_settings()


class UploadService:
    @staticmethod
    def save_file(file: UploadFile) -> tuple[str, str]:
        upload_dir = settings.UPLOAD_DIR
        os.makedirs(upload_dir, exist_ok=True)
        file_path = os.path.join(upload_dir, f"{generate_id('file')}_{file.filename}")
        with open(file_path, "wb") as f:
            f.write(file.file.read())
        return file_path, file.filename

    @staticmethod
    def calculate_checksum(file_path: str) -> str:
        sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                sha256.update(chunk)
        return sha256.hexdigest()

    @staticmethod
    def parse_txt(content: str) -> list[dict]:
        lines = content.replace("\r\n", "\n").split("\n")
        pattern = r"^\s*(第[一二三四五六七八九十百千万0-9]+[章节回幕]|chapter\s+\d+)"
        chapters = []
        current = {"title": "第一章", "content": []}
        for line in lines:
            if re.match(pattern, line.strip(), re.IGNORECASE):
                if current["content"]:
                    chapters.append(current)
                current = {"title": line.strip(), "content": []}
            else:
                current["content"].append(line)
        if current["content"]:
            chapters.append(current)
        return chapters

    @staticmethod
    def parse_docx(file_path: str) -> list[dict]:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        content = "\n".join(p.text for p in doc.paragraphs)
        return UploadService.parse_txt(content)

    @staticmethod
    def upload(db: Session, project_id: str, file: UploadFile) -> dict:
        project = db.query(Project).filter(Project.id == project_id).first()
        if not project:
            raise NotFoundError("Project")

        file_ext = file.filename.split(".")[-1].lower()
        if file_ext not in ("txt", "docx"):
            raise BadRequestError("Only TXT and DOCX files are supported")

        file_path, filename = UploadService.save_file(file)
        checksum = UploadService.calculate_checksum(file_path)

        if file_ext == "txt":
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            chapters = UploadService.parse_txt(content)
            parser_version = "txt-parser-1.0"
        else:
            chapters = UploadService.parse_docx(file_path)
            parser_version = "docx-parser-1.0"

        # Remove existing source document and chapters
        existing = db.query(SourceDocument).filter(SourceDocument.project_id == project_id).first()
        if existing:
            db.delete(existing)
            db.query(Chapter).filter(Chapter.project_id == project_id).delete()
            db.commit()

        source_doc = SourceDocument(
            id=generate_id("doc"),
            project_id=project_id,
            filename=filename,
            file_type=file_ext,
            language="zh-CN",
            checksum=checksum,
            file_path=file_path,
            chapter_count=len(chapters),
            total_characters=sum(len(c["title"] + "\n".join(c["content"])) for c in chapters),
            parser_version=parser_version,
        )
        db.add(source_doc)

        for idx, ch in enumerate(chapters, 1):
            ch_content = "\n".join(ch["content"])
            db.add(Chapter(
                id=generate_id("chapter"),
                project_id=project_id,
                index=idx,
                title=ch["title"],
                summary=ch_content[:100] + "..." if len(ch_content) > 100 else ch_content,
                word_count=len(ch_content),
                content=ch_content,
            ))

        project.status = "processing"
        db.commit()
        logger.info(f"File uploaded for project {project_id}: {filename}, {len(chapters)} chapters")
        return {"message": "File uploaded and parsed successfully", "chapter_count": len(chapters)}
